# --- MANUAL TEMPLATE ---
# Import libraries
from docx import Document
from docx.shared import Inches
import openai
import credentials

#%%
# --- GPT3 PROMPT  ---
openai.api_key = credentials.OpenAI_Key

# header = input("What is the header of the document?")
# documentation = input("What do you want to document?")

header = "Assignment2: MongoDB and Python Flask Application"
documentation = "Write a documentation for a python flask application with a MongoDB database."


# Run GPT 3
response = openai.Completion.create(
    engine="text-davinci-003",
    prompt=documentation,
    temperature=0.2,
    max_tokens=400,
    top_p=1,
    frequency_penalty=0,
    presence_penalty=0.6
)

answer = response['choices'][0]['text']

print(answer)


#%%
# --- CREATE DOCUMENT ---
# Open template document `hello_world.docx`
document = Document('template.docx')

# Clear document
document._body.clear_content()

# Add title
document.add_heading(header, 1)

# Add paragraph
p = document.add_paragraph(answer)

# Save document
document.save('doc2.docx')